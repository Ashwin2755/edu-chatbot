import os
import shutil
from pathlib import Path
from typing import List, Optional
import uuid

# PDF and Docx processing
import PyPDF2
import docx

# LangChain and RAG
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.docstore.document import Document as LCDocument

import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # Optimized for educational content
            chunk_overlap=100  # Better context preservation
        )
        # Using OpenAI-compatible embeddings (could use local too, but let's stick to OpenAI for now if API key exists)
        # Fallback to local if no API key
        api_key = os.getenv("OPENROUTER_API_KEY")
        if api_key:
             # We use OpenAIEmbeddings but point it to a model that works well or use standard OpenAI ones
             self.embeddings = OpenAIEmbeddings(openai_api_key=api_key)
        else:
            # Fallback to a local embedding model if no API key (requires sentence-transformers)
            from langchain_community.embeddings import HuggingFaceEmbeddings
            self.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
            
        self.vector_db = None
        self._init_vector_db()

    def _init_vector_db(self):
        try:
            self.vector_db = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings
            )
            logger.info(f"✅ ChromaDB initialized at {self.persist_directory}")
        except Exception as e:
            logger.error(f"❌ Failed to initialize ChromaDB: {e}")

    def extract_text(self, file_path: str) -> str:
        suffix = Path(file_path).suffix.lower()
        text = ""
        
        try:
            if suffix == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            elif suffix == '.docx':
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif suffix in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            else:
                logger.warning(f"Unsupported file type: {suffix}")
                
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    async def add_document(self, file_path: str, filename: str, doc_id: str):
        text = self.extract_text(file_path)
        if not text:
            return None
            
        chunks = self.text_splitter.split_text(text)
        documents = [
            LCDocument(
                page_content=chunk,
                metadata={"source": filename, "doc_id": doc_id}
            ) for chunk in chunks
        ]
        
        if self.vector_db:
            self.vector_db.add_documents(documents)
            # No need to call persist() in newer versions of ChromaDB, but good for safety
            # self.vector_db.persist()
            return text
        return None

    def get_relevant_context(self, query: str, k: int = 3) -> str:
        if not self.vector_db:
            return ""
            
        try:
            results = self.vector_db.similarity_search(query, k=k)
            context = "\n\n".join([doc.page_content for doc in results])
            return context
        except Exception as e:
            logger.error(f"Error retrieving context: {e}")
            return ""

    def clear_database(self):
        if os.path.exists(self.persist_directory):
            shutil.rmtree(self.persist_directory)
        self._init_vector_db()
